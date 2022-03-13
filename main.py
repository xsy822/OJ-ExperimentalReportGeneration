from docx.oxml.ns import qn
from docx.shared import Pt
from docx import Document
import os
import json
import requests


def getCodes():
    '''获取提交代码'''
    # 获取配置
    with open('user.json', 'r', encoding='utf-8') as f:
        userData = json.loads(f.read())
    headers = {'Cookie': userData['cookie']}
    url = userData['url']
    url1 = userData['url1']
    # 分页获取所有提交记录
    result = []
    maxPage = 1
    nowPage = 1
    while True:
        httpGet = requests.get(url[:77]+str(nowPage)+url[78:]+str((nowPage-1)*100),
                               headers=headers)
        result.extend(json.loads(httpGet.text)['data']['results'])
        data = json.loads(httpGet.text)['data']
        if nowPage == 1 and data['total'] > 100:
            maxPage = data['total']//100+1
        if nowPage == maxPage:
            break
        nowPage += 1
    problems = {}
    for i in result:
        if i['result'] == 0:
            problems[i['problem']] = i['id']
    ans = {}
    for key, value in problems.items():
        httpGet = requests.get(url1+value, headers=headers)
        ans[int(key) % 100+1000] = json.loads(httpGet.text)['data']['code']
    return ans


def lsDir():
    '''列举模板problem'''
    ans = {}
    ls = os.listdir('模板')
    for i in ls:
        ans[i] = list(map(lambda x: int(x[-9:-5]),
                          filter(lambda x: not x[0] == '实', os.listdir('模板/'+i))))
    return ans


def change(url, problem, code):
    '''修改word'''
    with open('user.json', 'r', encoding='utf-8') as f:
        userData = json.loads(f.read())
    #'AC-学号-姓名- 题号'
    document = Document('模板/{}/AC-学号-姓名- 题号{}.docx'.format(url, problem))
    table = document.tables[0]
    # 修改表格
    run = table.cell(1, 1).paragraphs[0].add_run(userData['班级'])
    font = run.font
    font.name = '微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    font.size = Pt(8)
    run = table.cell(1, 3).paragraphs[0].add_run(userData['姓名'])
    font = run.font
    font.name = '微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    font.size = Pt(8)
    run = table.cell(1, 5).paragraphs[0].add_run(userData['学号'])
    font = run.font
    font.name = '微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    font.size = Pt(8)
    # 贴代码
    p = document.paragraphs
    flag = 0
    for i in p:
        if flag:
            run = i.add_run(code)
            font = run.font
            font.name = '微软雅黑'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            font.size = Pt(8)
            break
        if i.text == '源代码：':
            flag = 1
    if not os.path.exists('模板/{}/{}已完成'.format(url, url)):
        os.mkdir('模板/{}/{}已完成'.format(url, url))
    document.save('模板/{}/{}已完成/AC-{}-{}- 题号{}.docx'.format(url,
                                                           url, userData['学号'], userData['姓名'], problem))


muBan = lsDir()
codes = getCodes()
finished = {}
for problem, code in codes.items():
    for i in muBan:
        if problem in muBan[i]:
            change(i, problem, code)
            if i in finished:
                finished[i].append(problem)
            else:
                finished[i] = [problem]
            muBan[i].remove(problem)
            if muBan[i] == []:
                print(i+'已全部完成')
print('此次修改了:')
for key, value in finished.items():
    print(key, value)
print('还未完成的：')
for key, value in muBan.items():
    print(key, value)
